from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

out = '/Users/Ezra/Desktop/Parah_Adumah_Hebrew_English_Side_by_Side.docx'

pages = [
    ('Page 1', [
        ('טהרת פרה אדומה', 'Purification of the Red Heifer'),
        ('ב"ה זאת חקת התורה אשר צוה ה\' לאמר דבר אל', 'Blessed be God. “This is the statute of the Torah which the Lord commanded, saying: Speak to”'),
        ('בני ישראל ויקחו אליך פרה אדומה תמימה וגו\'.', '“the Children of Israel, and let them take to you a perfect red heifer, etc.”'),
        ('שאמנם חוק הוא ואין להשיב הטעם אך י"ל ענינה', 'Indeed, it is a statute whose reason cannot be explained, but its meaning can be said to be'),
        ('ע"ד העבודה. וברש"י מביא מיסודו של ר\' משה', 'in the aspect of divine service. And Rashi brings, from Rabbi Moshe'),
        ('הדרשן, משל לבן שפחה שטינף פלטין, אמרו תבוא', 'HaDarshan, a parable of a maidservant’s son who dirtied a palace: they said, let the mother come'),
        ('יש לבאר ענין חוקת הפרה המטהרת טמאים.', 'to explain the matter of the statute of the heifer that purifies the impure.'),
    ]),
    ('Page 2', [
        ('חיקת', 'Chukat'),
        ('אמר הקב"ה צאו בנה, תביאו פרה ותכפרו על', 'Said the Holy One, Blessed Be He: go bring a cow and atone for'),
        ('העגל. וצ"ב דלכאורה הרי דיני טומאה וטהרה ומה', 'the calf. But this needs explanation, because seemingly the laws of impurity and purity and what'),
        ('ענינם אצל כפרה. עוד יש לעמוד על הדיוקים', 'do they have to do with atonement? Also, one must examine the precise wording'),
        ('בפרשה, מה שהעידו המפרשים על לשון הכתוב וידבר', 'in the portion, what the commentators noted about the verse “And the Lord spoke”'),
        ('ה\' אל משה ואהרן לאמר זאת חקת התורה אשר', 'to Moses and Aaron, saying: “This is the statute of the Torah that'),
        ('צוה ה\' לאמר וגו\'. שנאמר כאן ב\' פעמים לאמר.', 'the Lord commanded, saying,” etc. Here the word “saying” appears twice.'),
        ('כן צ"ב מד"כ יקחו אליך, דפרש"י לעולם היא', 'Also, it needs explanation from the phrase “they shall take to you,” for Rashi says it is always'),
        ('נקראת על שם פרה מעשה משה, מדוע מצות פרה', 'called after the red heifer, Moses’ deed; why does the commandment of the red'),
        ('אדומה במיחד מתייחסת אל משה, והלא כל המצוות', 'heifer especially relate to Moses, when all commandments'),
        ('ניתנו ע"י משה רבינו ולמה דווקא מצות פרה נקראת', 'were given through Moses our teacher, and why specifically is the commandment of the heifer called'),
        ('על שמו של משה. וכעין זה צ"ב בהא דהתורה', 'by Moses’ name? Similarly, it needs explanation why the Torah'),
        ('בכללותה מתייחסת על משה ונקראת על שמו, כמד"כ', 'in general relates to Moses and is called by his name, as it says'),
        ('זכרו תורת משה, והלא התורה ניתנה מאת ה\' ומשה', '“Remember the Torah of Moses,” when the Torah was given by the Lord and Moses'),
        ('היה רק שליח ומדוע היא מתייחסת על שמו.', 'was only a messenger, so why is it attributed to his name?'),
    ]),
    ('Page 3', [
        ('תיבות', 'Words / Terms'),
        ('האנוכיות והישות. ועד כדי כך מרחיקה הישות את', 'the ego and selfhood. And to such an extent does selfhood distance'),
        ('האדם מהקב"ה, כדאיתא בגמ\' (סוטה ה.) שהקב"ה', 'a person from the Holy One, Blessed Be He, as it is stated in the Gemara (Sotah 5a) that the Holy One, Blessed Be He'),
        ('אומר על המתגאה אין אני והוא יכולים לדור כאחד,', 'says about the arrogant person: “I and he cannot dwell together,”'),
        ('וכן אמרו חכמאתנו כאילו עובד ע"ז. ואם כל התורה', 'and our sages said it is as if he worships idolatry. And if all Torah'),
        ('והמצוות הם עצות להביא לדבקות בה\'', 'and commandments are advice to bring one to cleaving to the Lord, then selfhood'),
        ('מרחיקה את האדם מקיום המצוות היסודיות של ואהבת', 'distances a person from fulfilling the foundational commandments of “And you shall love'),
        ('ה\' אלקיך בכל לבבך ובכל נפשך ובכל מאדך,', 'the Lord your God with all your heart and with all your soul and with all your might,”'),
        ('כדאיתא בספה"ק שאבת ישראל סביבה לאהבת ה\'', 'as it is stated in the holy books that the love of Israel surrounds the love of God,'),
        ('ומי שאינו אוהב את הזולת אינו יכול לאהוב את', 'and one who does not love others cannot love'),
        ('השי"ת. והיינו כי מי שבכל עניניו נתונים להשביע', 'the Holy One, Blessed Be He. This is because a person whose affairs are all devoted to satisfying'),
        ('את רצון עצמו שוב אינו פנוי לאהבת הזולת ואהבת', 'his own will is no longer available to love others and to love'),
        ('השי"ת. והרי אלו ב\' הפכים המנגדים זה לזה, התורה', 'the Holy One, Blessed Be He. These are two opposites that contradict each other: the Torah'),
        ('והפצוות תכליתם להביא את יהודי לדבקות בה\'', 'and commandments aim to bring a Jew to cleaving to Him,'),
        ('ואילו האנוכיות והישות מרחיקות אותו מדבקות בה\'.', 'whereas ego and selfhood distance him from cleaving to Him.'),
    ]),
    ('Page 4', [
        ('קי תשובות', 'Responsum 100'),
        ('למדתו של משה רבינו לספור את הישות והחומריות.', 'Our teacher Moses taught the nullification of ego and materialism.'),
        ('ומזה"פ נתי התורה בכללותה נקראת תורת משה,', 'And from this, the Torah in its entirety is called the Torah of Moses,'),
        ('כי אין דברי תורה מתקיימין אלא במי שממית עצמו', 'because the words of Torah are sustained only by one who “kills” himself over it,'),
        ('עליה, שמבטל את כל ישותו ועצמיותו.', 'meaning one who nullifies all his ego and selfhood.'),
        ('ופ"פ האמור יבואר גם מאמר הבעש"ט הק\'', 'And according to this, the saying of the holy Baal Shem Tov is also explained'),
        ('על מאה"ל מת בערב שבת סימן יפה לו, שהקשה', 'regarding “one who dies on Friday evening — it is a good sign for him,” which he asked:'),
        ('וכי יש לאדם ברירה אימתי למות ומאי קמ"ל בזה.', 'does a person have a choice as to when he dies, and what is this teaching us?'),
        ('ופירש הכוונה, דיהודי העומד בשער השבת ואינו', 'And he explained that a Jew standing at the threshold of Shabbat and not'),
        ('מוצא במקש כמה להכנס לשבת, שהרי שבת ענינה', 'finding a way to enter Shabbat, since the essence of Shabbat'),
        ('דבקות בה\'', 'is cleaving to the Lord, while he feels himself as a “self,” which is the opposite of'),
        ('הוא מרגיש את עצמו יש, שהוא היפך', 'cleaving, where “I and He cannot dwell together,” how can he enter'),
        ('הדבקות שאין אני והוא יכולין ללמוד, ואיך יכנס', 'Shabbat? His advice is to “die” on Friday evening, that is, to nullify himself'),
        ('לשבת. עצתו לקיים מה בערב שבת, שיבטל עצמו', 'completely in the aspect of nothingness in order to enter Shabbat. This is a good sign'),
        ('לגמרי בבחי\' אין כמה כדי להיכנס לשבת, וזהו סימן', 'for him to enter through the holiness of Shabbat, meaning that through nullification'),
        ('יפה לו להיכנס על ידי הש"ק, והיינו שע"י ביטול', 'of ego in the aspect of “dying on Friday evening,” a Jew reaches cleaving'),
        ('הישות בבחי\' מת בערב שבת מגיע יהודי לדבקות', 'to the Lord, which is the purpose of Torah and commandments.'),
    ]),
]

# Fix obviously broken Hebrew/English balance by using a clean continuation for page 4 last lines.
pages[3][1].extend([
    ('בה\' שהיא תכלית התורה והמצוות, ושם בפרש"ב', 'to the Lord, which is the purpose of Torah and commandments, and there in the portion of'),
    ('(שמות לא) כי אות היא ביני וביניכם, שגם בשבת', '“for it is a sign between Me and you,” that even on Shabbat'),
    ('יש את הכח להביא לדבקות כזו כמו כח התורה.', 'there is power to bring about such cleaving, like the power of Torah.'),
    ('וזהו ענין שבעת ימי המיתה, שבכך י\' ימים יש', 'And this is the meaning of the seven days of mourning, that in these days there is'),
    ('תמיד שבת, והשבת היא בבחי\' ובזים השביעי ישר,', 'always Shabbat.'),
    ('שמביאה יהודי לעקור את הישות ולחיות רבוק בה\'', 'which brings a Jew to uproot ego and live cleaving to the Lord,'),
    ('שהדרך לעקור את הישות היא ע"י יסוף את הפה,', 'and the way to uproot ego is through...'),
    ('ע"י אש של קדושה, אש של תורה ואש של שבת', 'by the fire of holiness, the fire of Torah, and the fire of Shabbat'),
    ('המביאים יהודי לתכלית של הדבקות בה\'.', 'which bring a Jew to the purpose of cleaving to the Lord.'),
])

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.6)
sec.bottom_margin = Inches(0.6)
sec.left_margin = Inches(0.6)
sec.right_margin = Inches(0.6)
for s in doc.styles:
    pass

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Parah Adumah — Hebrew / English Side-by-Side Translation')
run.bold = True
run.font.size = Pt(16)

for page_title, rows in pages:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(page_title)
    r.bold = True
    r.font.size = Pt(12)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Hebrew'
    hdr[1].text = 'English'
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for h, e in rows:
        cells = table.add_row().cells
        cells[0].text = h
        cells[1].text = e
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_page_break()

doc.save(out)
print(out)
